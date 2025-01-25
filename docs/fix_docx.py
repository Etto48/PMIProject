import docx
import os
import subprocess as sp

import docx.styles
import docx.styles.styles

def fix_table_style(docx_file):
    doc = docx.Document(docx_file)
    # for style in doc.styles:
    #     print(style.name)
    for table in doc.tables:
        table.autofit = True
        # for row in table.rows:
        #     for cell in row.cells:
        #         cell.paragraphs[0].style = None
    doc.save(docx_file)

if __name__ == '__main__':
    template_path = os.path.join(os.path.dirname(__file__), 'main.docx')
    fix_table_style(template_path)