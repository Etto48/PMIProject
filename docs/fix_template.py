import docx
import os
import subprocess as sp

def fix_table_style(docx_file):
    doc = docx.Document(docx_file)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.paragraphs[0].style = None
    doc.save(docx_file)

if __name__ == '__main__':
    template_path = os.path.join(os.path.dirname(__file__), 'main.docx')
    fix_table_style(template_path)